from __future__ import annotations

import re
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "BUILD_PLAN.md"
OUTPUT = ROOT / "github_analytics_pipeline_build_plan.docx"
ARTIFACTS = ROOT / ".artifacts"
ARCHITECTURE_IMAGE = ARTIFACTS / "architecture-overview.png"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

# Named compactness override for the compact_reference_guide preset: use
# 0.78-inch vertical margins, 10.5-point body text, and 1.18 line spacing.
# Horizontal margins and table geometry retain the preset's required values.
COMPACT_VERTICAL_MARGIN_INCHES = 0.78
COMPACT_BODY_SIZE_PT = 10.5
COMPACT_BODY_LINE_SPACING = 1.18
FIXED_ZIP_TIMESTAMP = (2000, 1, 1, 0, 0, 0)

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "1F2937"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_formatted_runs(paragraph, text: str, size: float | None = None, color: str = BLACK) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run, size=size, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Cascadia Mono", size=(size or 11) - 0.5, color=DARK_BLUE)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_element = table._tbl
    table_properties = table_element.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = table_properties.find(qn(tag))
        if existing is not None:
            table_properties.remove(existing)

    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_properties.extend([table_width, table_indent, layout])

    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths_dxa[index]))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def choose_table_widths(headers: list[str]) -> list[int]:
    normalized = [header.lower() for header in headers]
    if normalized == ["topic", "key", "purpose", "local partitions"]:
        return [2500, 1250, 4410, 1200]
    if normalized == ["column", "purpose"]:
        return [2300, 7060]
    if normalized == ["dag", "schedule", "responsibilities"]:
        return [2100, 1700, 5560]
    if normalized == ["layer", "representative models", "responsibility"]:
        return [1450, 3600, 4310]
    if normalized == ["metric", "mvp status", "evidence"]:
        return [2200, 2300, 4860]
    if normalized == ["technology", "mvp choice", "rationale"]:
        return [1750, 2550, 5060]
    if normalized == ["technology", "decision trigger"]:
        return [2600, 6760]
    if normalized == ["day", "focus", "exit evidence"]:
        return [650, 2900, 5810]
    column_count = len(headers)
    base = CONTENT_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(document: Document, rows: list[list[str]]) -> None:
    widths = choose_table_widths(rows[0])
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_formatted_runs(paragraph, value, size=8.6 if len(rows[0]) >= 3 else 9)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)


def add_paragraph_shading(paragraph, fill: str, border: str | None = None) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    paragraph_properties.append(shading)
    if border:
        borders = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            edge = OxmlElement(f"w:{side}")
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "4")
            edge.set(qn("w:space"), "6")
            edge.set(qn("w:color"), border)
            borders.append(edge)
        paragraph_properties.append(borders)


def add_callout(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.18
    add_paragraph_shading(paragraph, CALLOUT, LIGHT_BLUE)
    add_formatted_runs(paragraph, text, size=10.2, color=NAVY)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    add_paragraph_shading(paragraph, LIGHT_GRAY, "D8DEE8")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, name="Cascadia Mono", size=7.8, color=DARK_BLUE)
        if index < len(lines) - 1:
            run.add_break()


def get_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_centered_text(draw, box, lines, font, fill) -> None:
    x0, y0, x1, y1 = box
    line_height = int(font.size * 1.3) if hasattr(font, "size") else 18
    total_height = line_height * len(lines)
    y = y0 + (y1 - y0 - total_height) / 2
    for line in lines:
        bounds = draw.textbbox((0, 0), line, font=font)
        width = bounds[2] - bounds[0]
        draw.text((x0 + (x1 - x0 - width) / 2, y), line, font=font, fill=fill)
        y += line_height


def draw_arrow(draw, start, end, color=BLUE, width=5) -> None:
    if not color.startswith("#"):
        color = f"#{color}"
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 16 * direction, ey - 10), (ex - 16 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - 16 * direction), (ex + 10, ey - 16 * direction)]
    draw.polygon(points, fill=color)


def build_architecture_image() -> None:
    ARTIFACTS.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 1030), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    title_font = get_font(44, bold=True)
    box_font = get_font(29, bold=True)
    small_font = get_font(22)
    draw.text((70, 35), "Validated event path, independent consumers, and batch recovery", font=title_font, fill=f"#{NAVY}")

    boxes = {
        "github": (70, 165, 360, 315),
        "api": (510, 165, 840, 315),
        "kafka": (990, 165, 1300, 315),
        "raw": (1450, 95, 1740, 245),
        "pr": (1450, 305, 1740, 455),
        "backfill": (70, 590, 410, 750),
        "airflow": (530, 590, 840, 750),
        "dbt": (990, 590, 1300, 750),
        "marts": (1450, 590, 1740, 750),
        "metabase": (1450, 840, 1740, 975),
    }
    labels = {
        "github": (["GitHub App", "webhooks"], "signed deliveries"),
        "api": (["FastAPI receiver"], "HMAC + envelope"),
        "kafka": (["Kafka raw topic"], "repo-keyed replay log"),
        "raw": (["Warehouse writer"], "append-only raw events"),
        "pr": (["PR monitor"], "state + alert outbox"),
        "backfill": (["GitHub APIs"], "GraphQL + REST history"),
        "airflow": (["Airflow"], "backfill + refresh"),
        "dbt": (["dbt"], "tested transformations"),
        "marts": (["Analytics marts"], "metrics + coverage"),
        "metabase": (["Metabase"], "three dashboards"),
    }

    for key, box in boxes.items():
        fill = LIGHT_BLUE if key not in {"kafka", "marts"} else "DCEAF7"
        draw.rounded_rectangle(box, radius=20, fill=f"#{fill}", outline=f"#{BLUE}", width=4)
        lines, subtitle = labels[key]
        draw_centered_text(draw, (box[0], box[1] - 15, box[2], box[3] - 10), lines, box_font, f"#{NAVY}")
        bounds = draw.textbbox((0, 0), subtitle, font=small_font)
        width = bounds[2] - bounds[0]
        draw.text(((box[0] + box[2] - width) / 2, box[3] - 44), subtitle, font=small_font, fill=f"#{MUTED}")

    draw_arrow(draw, (360, 240), (510, 240))
    draw_arrow(draw, (840, 240), (990, 240))
    draw_arrow(draw, (1300, 205), (1450, 170))
    draw_arrow(draw, (1300, 275), (1450, 380))
    draw_arrow(draw, (410, 670), (530, 670))
    draw_arrow(draw, (840, 670), (990, 670))
    draw_arrow(draw, (1300, 670), (1450, 670))
    draw_arrow(draw, (1595, 750), (1595, 840))

    draw.line([(240, 590), (240, 505), (1595, 505), (1595, 245)], fill=f"#{BLUE}", width=5)
    draw.polygon([(1595, 245), (1585, 265), (1605, 265)], fill=f"#{BLUE}")
    draw.text((695, 470), "idempotent backfill writes to the same raw contract", font=small_font, fill=f"#{DARK_BLUE}")

    draw.line([(1595, 455), (1595, 530), (1100, 530), (1100, 590)], fill=f"#{BLUE}", width=5)
    draw.polygon([(1100, 590), (1090, 570), (1110, 570)], fill=f"#{BLUE}")

    image.save(ARCHITECTURE_IMAGE)


def add_numbering(document: Document) -> dict[str, int]:
    numbering = document.part.numbering_part.element
    existing_abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abstract = max(existing_abstract_ids, default=0) + 1
    next_num = max(existing_num_ids, default=0) + 1
    result = {}
    new_abstracts = []
    new_numbers = []

    definitions = [
        ("bullet", "bullet", "-", "Calibri"),
        ("decimal", "decimal", "%1.", "Calibri"),
        ("check", "bullet", "\u2610", "Segoe UI Symbol"),
    ]
    for offset, (name, number_format, level_text, font_name) in enumerate(definitions):
        abstract_id = next_abstract + offset
        num_id = next_num + offset
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), number_format)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), level_text)
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        paragraph_properties = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "540")
        indent.set(qn("w:hanging"), "270")
        paragraph_properties.extend([tabs, indent])
        run_properties = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font_name)
        fonts.set(qn("w:hAnsi"), font_name)
        run_properties.append(fonts)
        level.extend([start, fmt, text, justification, paragraph_properties, run_properties])
        abstract.append(level)
        new_abstracts.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_reference = OxmlElement("w:abstractNumId")
        abstract_reference.set(qn("w:val"), str(abstract_id))
        num.append(abstract_reference)
        new_numbers.append(num)
        result[name] = num_id

    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    for offset, abstract in enumerate(new_abstracts):
        numbering.insert(first_num_index + offset, abstract)
    for num in new_numbers:
        numbering.append(num)
    return result


def apply_numbering(paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_properties.extend([level, number])
    paragraph_properties.append(num_properties)


def clone_numbering_instance(document: Document, source_num_id: int) -> int:
    numbering = document.part.numbering_part.element
    source = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == source_num_id
    )
    abstract_id = source.find(qn("w:abstractNumId")).get(qn("w:val"))
    next_num_id = max(
        (int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))),
        default=0,
    ) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), abstract_id)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.extend([abstract_reference, level_override])
    numbering.append(num)
    return next_num_id


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(COMPACT_VERTICAL_MARGIN_INCHES)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(COMPACT_VERTICAL_MARGIN_INCHES)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(COMPACT_BODY_SIZE_PT)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = COMPACT_BODY_LINE_SPACING
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(COMPACT_BODY_SIZE_PT)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = COMPACT_BODY_LINE_SPACING

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    run = header_paragraph.add_run("GITHUB ENGINEERING ANALYTICS  |  BUILD PLAN")
    set_run_font(run, size=8.5, color=MUTED)
    run.bold = True

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    label = footer_paragraph.add_run("Page ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(footer_paragraph)


def add_title_block(document: Document) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("TECHNICAL BUILD PLAN")
    set_run_font(run, size=9.5, color=BLUE)
    run.bold = True

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("GitHub Engineering Analytics Pipeline")
    set_run_font(run, size=23, color=NAVY)
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Architecture Review and 15-Day Implementation Plan")
    set_run_font(run, size=13, color=MUTED)

    metadata = [
        ("Status", "Ready for implementation"),
        ("Reviewed", "July 21, 2026"),
        ("Target", "Portfolio-grade MVP in 15 focused working days"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10, color=BLACK)
        label_run.bold = True
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10, color=BLACK)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(5)
    rule.paragraph_format.space_after = Pt(8)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    rule._p.get_or_add_pPr().append(borders)


def normalize_docx_archive(path: Path) -> None:
    """Rewrite the DOCX package with stable member timestamps and ordering."""
    with ZipFile(path, "r") as source_archive:
        members = [(info, source_archive.read(info.filename)) for info in source_archive.infolist()]
        archive_comment = source_archive.comment

    with tempfile.NamedTemporaryFile(dir=path.parent, suffix=".docx", delete=False) as temporary_file:
        temporary_path = Path(temporary_file.name)

    try:
        with ZipFile(temporary_path, "w", compression=ZIP_DEFLATED, compresslevel=9) as output_archive:
            output_archive.comment = archive_comment
            for source_info, payload in members:
                normalized_info = ZipInfo(source_info.filename, FIXED_ZIP_TIMESTAMP)
                normalized_info.compress_type = ZIP_DEFLATED
                normalized_info.create_system = source_info.create_system
                normalized_info.external_attr = source_info.external_attr
                normalized_info.internal_attr = source_info.internal_attr
                normalized_info.comment = source_info.comment
                output_archive.writestr(
                    normalized_info,
                    payload,
                    compress_type=ZIP_DEFLATED,
                    compresslevel=9,
                )
        temporary_path.replace(path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def parse_markdown(document: Document, numbering_ids: dict[str, int]) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 1."))
    lines = lines[start:]
    index = 0
    current_section = ""
    current_decimal_num_id = None
    used_base_decimal_numbering = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if not numbered:
            current_decimal_num_id = None
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            block = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            if language == "mermaid":
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(3)
                run = paragraph.add_run()
                run.add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.35))
                doc_properties = run._r.xpath(".//wp:docPr")
                if doc_properties:
                    doc_properties[0].set("descr", "Architecture flow from GitHub webhooks through Kafka consumers, PostgreSQL, Airflow, dbt, and Metabase")
                    doc_properties[0].set("title", "Revised architecture overview")
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                caption_run = caption.add_run("Figure 1. Revised MVP architecture and durable processing boundaries")
                set_run_font(caption_run, size=8.5, color=MUTED)
                caption_run.italic = True
            else:
                add_code_block(document, block)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = []
            for table_line in table_lines:
                cells = [cell.strip() for cell in table_line.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                rows.append(cells)
            add_table(document, rows)
            continue

        if stripped.startswith("## "):
            current_section = stripped[3:]
            document.add_paragraph(current_section, style="Heading 1")
            index += 1
            continue
        if stripped.startswith("### "):
            document.add_paragraph(stripped[4:], style="Heading 2")
            index += 1
            continue
        if stripped.startswith("#### "):
            document.add_paragraph(stripped[5:], style="Heading 3")
            index += 1
            continue

        if stripped.startswith("> "):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith("> "):
                quote_lines.append(lines[index].strip()[2:])
                index += 1
            add_callout(document, " ".join(quote_lines))
            continue

        if re.match(r"^- \[[ xX]\] ", stripped):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15
            apply_numbering(paragraph, numbering_ids["check"])
            add_formatted_runs(paragraph, re.sub(r"^- \[[ xX]\] ", "", stripped), size=10.2)
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph()
            source_list = current_section.startswith("15. Primary research basis")
            paragraph.paragraph_format.space_after = Pt(2 if source_list else 4)
            paragraph.paragraph_format.line_spacing = 1.05 if source_list else 1.18
            apply_numbering(paragraph, numbering_ids["bullet"])
            add_formatted_runs(paragraph, stripped[2:], size=9.2 if source_list else None)
            index += 1
            continue

        if numbered:
            if current_decimal_num_id is None:
                if used_base_decimal_numbering:
                    current_decimal_num_id = clone_numbering_instance(document, numbering_ids["decimal"])
                else:
                    current_decimal_num_id = numbering_ids["decimal"]
                    used_base_decimal_numbering = True
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.18
            apply_numbering(paragraph, current_decimal_num_id)
            add_formatted_runs(paragraph, numbered.group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith(("#", "- ", "> ", "|", "```"))
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        add_formatted_runs(paragraph, " ".join(paragraph_lines))


def main() -> None:
    build_architecture_image()
    document = Document()
    configure_document(document)
    numbering_ids = add_numbering(document)
    add_title_block(document)
    parse_markdown(document, numbering_ids)

    document.core_properties.title = "GitHub Engineering Analytics Pipeline - Architecture Review and 15-Day Build Plan"
    document.core_properties.subject = "Portfolio-grade data engineering project plan"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.keywords = "GitHub, Kafka, Airflow, dbt, PostgreSQL, engineering analytics"
    document.core_properties.comments = "Generated from BUILD_PLAN.md"

    document.save(OUTPUT)
    normalize_docx_archive(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
